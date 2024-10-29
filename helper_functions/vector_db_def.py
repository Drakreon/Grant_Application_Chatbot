import os
import tiktoken
import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.retrievers import ParentDocumentRetriever
from langchain.storage import InMemoryStore
# from langchain.chains import RetrievalQA
#from langchain_experimental.text_splitter import SemanticChunker
#from langchain.retrievers.multi_query import MultiQueryRetriever

###### RUN CODE [ONCE] TO GET VECTOR_DB #########

# CONSTANTS
EMBEDDINGS_MODEL_NAME = 'text-embedding-3-small'
LLM_MODEL_NAME = "gpt-4o-mini"

# INITIALISE CREDENTIALS
if load_dotenv('.env'):
   # for local development
   OPENAI_KEY = os.getenv('OPENAI_API_KEY')
else:
   OPENAI_KEY = st.secrets['OPENAI_API_KEY']

# INITIALISE MODELS
embeddings_model = OpenAIEmbeddings(model=EMBEDDINGS_MODEL_NAME)
llm=ChatOpenAI(temperature=0, model=LLM_MODEL_NAME)
client = OpenAI(api_key=OPENAI_KEY)

def get_embedding(input, model='text-embedding-3-small'):
    response = client.Embedding.create(
        input=input,
        model=model
    )
    return [x.embedding for x in response.data]

def get_completion(prompt, model=LLM_MODEL_NAME, temperature=0, top_p=1.0, max_tokens=256, n=1, json_output=False):
    if json_output == True:
      output_json_structure = {"type": "json_object"}
    else:
      output_json_structure = None

    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        n=1,
        response_format=output_json_structure,
    )
    return response.choices[0].message.content

def count_tokens(text):
    encoding = tiktoken.encoding_for_model(LLM_MODEL_NAME)
    return len(encoding.encode(text))

def count_tokens_from_message_rough(messages):
    encoding = tiktoken.encoding_for_model(LLM_MODEL_NAME)
    value = ' '.join([x.get('content') for x in messages])
    return len(encoding.encode(value))

folder_path = r"C:/Users/Sin Kuan Tan/OneDrive/Coding/Streamlit/Grant_Application_Chatbot/data"

documents = os.listdir(folder_path)
all_documents_content = []

# def read_word_doc(file_path):
#     doc = Document(file_path)

#     # Extract paragraphs text
#     doc_text = "\n".join([para.text for para in doc.paragraphs])

#     # Extract table data
#     table_data = []
#     for table in doc.tables:
#         for row in table.rows:
#             row_data = [cell.text.strip() for cell in row.cells]
#             table_data.append(row_data)

#     return doc_text, table_data

# def read_csv_file(file_path):
#     df = pd.read_csv(file_path)
#     return None, df.to_string()  

# Function to read a PDF file (text + attempt to extract tables)
def read_pdf_file(file_path):
    reader = PdfReader(file_path)

    # Extract general text from PDF
    text = ""
    for page in reader.pages:
        text += page.extract_text()

    # Placeholder for extracting table data from PDF (requires advanced tools)
    table_data = []  # More advanced libraries like `pdfplumber` can be used here

    return text, table_data

# Process each file based on its extension and prepare Document objects
for doc in documents:
    file_path = os.path.join(folder_path, doc)

    if doc.endswith('.docx'):
        # # Read Word file (text and tables)
        # content_text, table_content = read_word_doc(file_path)
        # # Combine text and table content
        # combined_content = content_text
        # if table_content:
    #     #     combined_content += "\n\n" + "\n".join([str(row) for row in table_content])

    #     # Create a LangChain Document object
    #     document_obj = LangchainDocument(page_content=combined_content, metadata={"filename": doc})
    #     all_documents_content.append(document_obj)

    # elif doc.endswith('.csv'):
    #     # Read CSV file (as table)
    #     content_text, table_content = read_csv_file(file_path)
    #     # Combine text and table content (CSV doesn't have text content)
    #     combined_content = table_content

    #     # Create a LangChain Document object
    #     document_obj = LangchainDocument(page_content=combined_content, metadata={"filename": doc})
    #     all_documents_content.append(document_obj)

    elif doc.endswith('.pdf'):
        # Read PDF file (text and attempt to get tables)
        content_text, table_content = read_pdf_file(file_path)
        # Combine text and table content
        combined_content = content_text
        if table_content:
            combined_content += "\n\n" + "\n".join([str(row) for row in table_content])

        # Create a LangChain Document object
        document_obj = LangchainDocument(page_content=combined_content, metadata={"filename": doc})
        all_documents_content.append(document_obj)

### NAIVE DB ####
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, length_function=count_tokens)
splitted_documents = text_splitter.split_documents(all_documents_content)
vectordb = Chroma.from_documents(documents=splitted_documents,embedding=embeddings_model,collection_name="naive_splitter", persist_directory="./vector_db")

# #### Parent and Child DB ####
# parent_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, length_function=count_tokens)
# child_splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=40, length_function=count_tokens)

# vectordb = Chroma(collection_name="parent_child", embedding_function=embeddings_model,persist_directory="./vector_db_PC")
# store = InMemoryStore()

# retriever = ParentDocumentRetriever(
#     vectorstore=vectordb,
#     docstore=store,
#     child_splitter=child_splitter,
#     parent_splitter=parent_splitter,
#     search_kwargs={'k': 4}
# )

# retriever.add_documents(all_documents_content)